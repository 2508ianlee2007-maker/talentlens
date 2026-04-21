import os
import re
import time
import tempfile
from pathlib import Path
from typing import Dict, List, Any

from pypdf import PdfReader
from docx import Document as DocxDocument
from groq import Groq
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings

QWEN_MODEL = "qwen/qwen3-32b"
FINAL_CHUNK_SIZE = 700
FINAL_CHUNK_OVERLAP = 175
FINAL_K = 6
SYSTEM_PROMPT = """You are an HR screening assistant.
Your task is to compare a candidate's resume with a job description.

Focus on:
1. Relevant technical skills
2. Relevant work experience
3. Overall suitability for the role

Return the answer in a clear and concise format with:
- Candidate Summary
- Matching Skills
- Missing / Weak Areas
- Suitability Score out of 10
"""

_CHAT_SYSTEM_PROMPT = """You are helping a student demo an AI-powered CV screening system.
Answer based on the provided job descriptions, CV text, and screening results.
Be clear, concise, and practical. If the context is insufficient, say so.
"""

_EMBEDDINGS = None


def get_embedding_model():
    global _EMBEDDINGS
    if _EMBEDDINGS is None:
        _EMBEDDINGS = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    return _EMBEDDINGS


def read_pdf_text(path: str) -> str:
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def read_txt_text(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def read_docx_text(path: str) -> str:
    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())


def read_uploaded_file(uploaded_file) -> str:
    suffix = Path(uploaded_file.name).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.getbuffer())
        tmp_path = tmp.name
    try:
        if suffix == ".pdf":
            return read_pdf_text(tmp_path)
        if suffix == ".docx":
            return read_docx_text(tmp_path)
        if suffix == ".txt":
            return read_txt_text(tmp_path)
        raise ValueError(f"Unsupported file type: {suffix}")
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def preprocess_text(text: str) -> str:
    if not isinstance(text, str):
        text = str(text)
    text = text.replace("\xa0", " ").replace("\n", " ").replace("\t", " ").lower()
    text = re.sub(r'(\+?\d[\d\-\(\) ]{7,}\d)', ' ', text)
    text = re.sub(r'[^a-z0-9/\+\#\-\.\s]', ' ', text)
    text = re.sub(r'[\.]{2,}', ' ', text)
    text = re.sub(r'[-]{2,}', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def extract_score(text: str):
    if not isinstance(text, str):
        return None
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL | re.IGNORECASE)
    cleaned = text.replace("**", "").replace("__", "").replace("*", "").replace("\r", "\n")
    patterns = [
        r"suitability\s*score\s*[:\-–]?\s*(\d+(?:\.\d+)?)\s*/\s*10",
        r"suitability\s*score\s*[:\-–]?\s*(\d+(?:\.\d+)?)\s*out\s*of\s*(?:10|ten)",
        r"suitability\s*score\s*[:\-–]?\s*(\d+(?:\.\d+)?)",
        r"(?:^|\n)\s*(?:overall\s+)?(?:final\s+)?(?:score|rating)\s*[:\-–]?\s*(\d+(?:\.\d+)?)\s*/\s*10",
        r"(?:^|\n)\s*(?:overall\s+)?(?:final\s+)?(?:score|rating)\s*[:\-–]?\s*(\d+(?:\.\d+)?)\s*out\s*of\s*(?:10|ten)",
        r"(?:^|\n)\s*(?:overall\s+)?(?:final\s+)?(?:score|rating)\s*[:\-–]?\s*(\d+(?:\.\d+)?)",
        r"[\(\[]\s*(\d+(?:\.\d+)?)\s*/\s*10\s*[\)\]]",
    ]
    for pattern in patterns:
        m = re.search(pattern, cleaned, re.IGNORECASE)
        if m:
            val = float(m.group(1))
            if 0 <= val <= 10:
                return val
    fallback = re.findall(r"(?:suitability|score|rating)[^0-9\n]{0,20}(\d+(?:\.\d+)?)", cleaned, re.IGNORECASE)
    for item in reversed(fallback):
        val = float(item)
        if 0 <= val <= 10:
            return val
    return None


def extract_section(text: str, names: List[str]) -> str:
    if not isinstance(text, str):
        return ""
    cleaned = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL | re.IGNORECASE).strip()
    escaped = [re.escape(n) for n in names]
    pattern = (
        r"(?is)(?:^|\n)\s*(%s)\s*:\s*(.*?)"
        r"(?=\n\s*(?:Candidate Summary|Matching Skills|Missing\s*/\s*Weak Areas"
        r"|Missing Areas|Weak Areas|Suitability Score|Overall Score|Final Score)\s*:|\Z)"
    ) % "|".join(escaped)
    m = re.search(pattern, cleaned)
    if not m:
        return ""
    return re.sub(r"\s+", " ", m.group(2)).strip(" -\n\t")


def fmt_name(filename: str) -> str:
    return Path(filename).stem.replace("_", " ").replace("-", " ")


def verdict_label(score):
    if score is None:
        return "Unknown", "v-weak"
    if score >= 8:
        return "Strong Match", "v-strong"
    if score >= 6:
        return "Good Match", "v-good"
    if score >= 4:
        return "Partial Match", "v-partial"
    return "Weak Match", "v-weak"


def score_color(score):
    if score is None:
        return "#94a3b8"
    if score >= 8:
        return "#0f766e"
    if score >= 6:
        return "#1d4ed8"
    if score >= 4:
        return "#b45309"
    return "#b91c1c"


def build_user_prompt(jd: str, name: str, resume: str) -> str:
    return (
        f"Job Description:\n{jd}\n\nCandidate Resume File Name:\n{name}\n\n"
        f"Candidate Resume Text:\n{resume}\n\nPlease evaluate this candidate.\n\n"
        "Format your answer as:\nCandidate Summary:\nMatching Skills:\n"
        "Missing / Weak Areas:\nSuitability Score:"
    )


def build_rag_prompt(jd: str, name: str, context: str) -> str:
    return (
        f"Job Description:\n{jd}\n\nCandidate Resume File Name:\n{name}\n\n"
        f"Most Relevant Resume Sections (retrieved via RAG):\n{context}\n\n"
        "Based on the retrieved resume sections above, please evaluate this candidate.\n\n"
        "Format your answer as:\nCandidate Summary:\nMatching Skills:\n"
        "Missing / Weak Areas:\nSuitability Score:"
    )


def load_and_preprocess_files(jd_uploads, cv_uploads):
    jd_files, raw_jd, cv_files, raw_cv = {}, {}, {}, {}
    errors = []

    for f in jd_uploads or []:
        try:
            raw = read_uploaded_file(f)
            raw_jd[f.name] = raw
            jd_files[f.name] = preprocess_text(raw)
        except Exception as e:
            errors.append(f"JD '{f.name}': {e}")

    for f in cv_uploads or []:
        try:
            raw = read_uploaded_file(f)
            raw_cv[f.name] = raw
            cv_files[f.name] = preprocess_text(raw)
        except Exception as e:
            errors.append(f"CV '{f.name}': {e}")

    return {
        "jd_files": jd_files,
        "raw_jd_texts": raw_jd,
        "cv_files": cv_files,
        "raw_cv_texts": raw_cv,
        "errors": errors,
    }


def _build_chunks(cv_files: Dict[str, str]):
    documents = [Document(page_content=text, metadata={"source": name}) for name, text in cv_files.items()]
    splitter = RecursiveCharacterTextSplitter(chunk_size=FINAL_CHUNK_SIZE, chunk_overlap=FINAL_CHUNK_OVERLAP)
    return splitter.split_documents(documents)


def _get_rag_context(chunks, emb, cv_name: str, query: str) -> str:
    cand_chunks = [c for c in chunks if c.metadata.get("source") == cv_name]
    if not cand_chunks:
        return ""
    vectorstore = FAISS.from_documents(cand_chunks, emb)
    retriever = vectorstore.as_retriever(search_kwargs={"k": min(FINAL_K, len(cand_chunks))})
    return "\n\n".join(d.page_content for d in retriever.invoke(query))


def _call_llm(client: Groq, prompt: str) -> str:
    response = client.chat.completions.create(
        model=QWEN_MODEL,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.1,
        max_tokens=1024,
    )
    return response.choices[0].message.content


def screen_candidates(
    jd_files: Dict[str, str],
    cv_files: Dict[str, str],
    groq_key: str,
    use_rag: bool = True,
    delay: int = 1,
    progress_callback=None,
) -> Dict[str, List[Dict[str, Any]]]:
    if not groq_key or not groq_key.strip():
        raise ValueError("Groq API key is required.")
    if not jd_files:
        raise ValueError("No job descriptions loaded.")
    if not cv_files:
        raise ValueError("No CV files loaded.")

    groq_client = Groq(api_key=groq_key.strip())
    chunks = _build_chunks(cv_files) if use_rag else []
    emb = get_embedding_model() if use_rag else None

    total_pairs = max(len(jd_files) * len(cv_files), 1)
    step = 0
    all_results = {}

    for jd_name, cleaned_jd in jd_files.items():
        jd_res = []
        for cv_name, cleaned_cv in cv_files.items():
            step += 1
            if progress_callback:
                progress_callback(step, total_pairs, jd_name, cv_name)

            if use_rag:
                ctx = _get_rag_context(chunks, emb, cv_name, cleaned_jd[:500])
                prompt = build_rag_prompt(cleaned_jd, cv_name, ctx)
            else:
                prompt = build_user_prompt(cleaned_jd, cv_name, cleaned_cv)

            try:
                output = _call_llm(groq_client, prompt)
            except Exception as e:
                output = f"Error: {e}"

            jd_res.append({
                "cv_name": cv_name,
                "score": extract_score(output),
                "output": output,
                "summary": extract_section(output, ["Candidate Summary"]),
                "skills": extract_section(output, ["Matching Skills"]),
                "gaps": extract_section(output, ["Missing / Weak Areas", "Missing Areas", "Weak Areas"]),
                "mode": "RAG" if use_rag else "Direct",
            })
            if delay:
                time.sleep(delay)

        jd_res.sort(key=lambda x: x["score"] if x["score"] is not None else -1, reverse=True)
        all_results[jd_name] = jd_res

    return all_results


def build_chat_context(raw_jd_texts, raw_cv_texts, results) -> str:
    ctx = ""
    if raw_jd_texts:
        ctx += "=== JOB DESCRIPTIONS ===\n"
        for name, text in raw_jd_texts.items():
            ctx += f"\n--- {name} ---\n{text[:2500]}\n"
    if raw_cv_texts:
        ctx += "\n=== CANDIDATE CVS ===\n"
        for name, text in raw_cv_texts.items():
            ctx += f"\n--- {name} ---\n{text[:2500]}\n"
    if results:
        ctx += "\n=== SCREENING RESULTS ===\n"
        for jd_name, rows in results.items():
            ctx += f"\nJD: {jd_name}\n"
            for row in rows:
                ctx += (
                    f"  {row['cv_name']}: {row['score']}/10 "
                    f"({verdict_label(row['score'])[0]}) — "
                    f"{(row.get('summary') or '')[:150]}\n"
                )
    return ctx.strip()


def ask_about_candidates(groq_key: str, context: str, question: str, chat_history=None) -> str:
    if not groq_key or not groq_key.strip():
        raise ValueError("Groq API key is required.")
    client = Groq(api_key=groq_key.strip())
    messages = [{"role": "system", "content": _CHAT_SYSTEM_PROMPT}]
    if context:
        messages.append({"role": "system", "content": f"Context:\n{context}"})
    for item in (chat_history or []):
        role = item.get("role")
        content = item.get("content", "")
        if role in {"user", "assistant"} and content:
            messages.append({"role": role, "content": content})
    messages.append({"role": "user", "content": question})
    response = client.chat.completions.create(
        model=QWEN_MODEL,
        messages=messages,
        temperature=0.2,
        max_tokens=900,
    )
    return response.choices[0].message.content
