import streamlit as st
import os, re, time, tempfile
from pathlib import Path

import pandas as pd
from pypdf import PdfReader
from docx import Document as DocxDocument
from groq import Groq

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings

# ─── PAGE CONFIG ──────────────────────────────────────────────────────────────
st.set_page_config(page_title="TalentLens", page_icon="🔍", layout="wide")

st.markdown("""
<style>
    /* ── Force light mode — prevents system/browser dark mode from bleeding in ── */
    :root { color-scheme: light !important; }
    html, body,
    [data-testid="stApp"],
    [data-testid="stAppViewContainer"],
    [data-testid="stMain"],
    .main, .block-container {
        background-color: #f8fafc !important;
        color: #1e293b !important;
    }
    /* All general text stays dark */
    p, span, div, label, h1, h2, h3, h4, h5, h6,
    [data-testid="stMarkdownContainer"],
    .stMarkdown, .element-container { color: #1e293b !important; }
    /* Widget labels */
    .stSelectbox label, .stMultiSelect label, .stFileUploader label,
    .stCheckbox label, .stRadio label, .stSlider label,
    .stNumberInput label, .stTextInput label, .stTextArea label { color: #1e293b !important; }
    /* Expanders */
    [data-testid="stExpander"] {
        background: #ffffff !important;
        border: 1px solid #dbe4f0 !important;
        border-radius: 10px !important;
    }
    [data-testid="stExpander"] summary { color: #1e293b !important; }
    /* Metrics */
    [data-testid="stMetricValue"],
    [data-testid="stMetricLabel"],
    [data-testid="stMetricDelta"] { color: #1e293b !important; }
    /* Selectbox */
    [data-baseweb="select"] * { color: #1e293b !important; background: #fff !important; }
    /* Chat messages */
    [data-testid="stChatMessage"] { background: #fff !important; color: #1e293b !important; }
    /* Download button text */
    [data-testid="stDownloadButton"] button { color: #1e293b !important; }
    /* Tab panels */
    [data-testid="stTabs"] { background: transparent !important; }

    /* ── Global ── */
    [data-testid="stAppViewContainer"] { background: #f8fafc; }
    [data-testid="stSidebar"] { background: #eef4ff !important; }
    [data-testid="stSidebar"] * { color: #1e293b !important; }

    /* ── Sidebar input / uploader fixes ── */
    [data-testid="stSidebar"] .stTextInput input {
        background: #ffffff !important;
        border: 1px solid #cbd5e1 !important;
        color: #0f172a !important;
        border-radius: 10px;
    }
    [data-testid="stSidebar"] .stTextInput input::placeholder {
        color: #94a3b8 !important;
        opacity: 1 !important;
    }
    [data-testid="stSidebar"] .stTextInput > div {
        background: transparent !important;
    }
    [data-testid="stSidebar"] .stButton > button {
        background: #b9d8ff !important;
        color: #1e3a8a !important;
        border: 1px solid #9ec5fe !important;
        border-radius: 10px;
        box-shadow: none !important;
    }
    [data-testid="stSidebar"] .stButton > button:hover {
        background: #9ec5fe !important;
        color: #1e3a8a !important;
    }

    /* ── Upload blocks ── */
    [data-testid="stFileUploader"] {
        background: #fdf2f8 !important;
        border: 1px solid #f5c2d7 !important;
        border-radius: 14px !important;
        padding: 10px !important;
    }
    [data-testid="stFileUploader"] section {
        background: #fff7fb !important;
        border-radius: 12px !important;
        border: 1px dashed #e9a8c3 !important;
    }
    [data-testid="stFileUploader"] button {
        background: #fde2e7 !important;
        color: #9d174d !important;
        border: 1px solid #f8b4c8 !important;
        border-radius: 10px !important;
    }
    [data-testid="stFileUploader"] button:hover {
        background: #fbcfe8 !important;
        color: #831843 !important;
    }

    /* ── Tighter sidebar spacing ── */
    [data-testid="stSidebar"] .element-container { margin-bottom: 0.55rem !important; }
    [data-testid="stSidebar"] .stDivider { margin: 0.9rem 0 !important; }

    /* ── Tabs ── */
    div[data-baseweb="tab-list"] { gap: 0.5rem; border-bottom: 2px solid #dbe4f0; }
    button[data-baseweb="tab"] {
        height: auto; white-space: nowrap;
        padding: 0.6rem 1.2rem; border-radius: 10px 10px 0 0 !important;
        font-weight: 500;
    }
    button[data-baseweb="tab"][aria-selected="true"] {
        background: #ffffff !important; border-bottom: 2px solid #7c9fd6 !important;
        color: #355c9a !important;
    }

    /* ── Cards ── */
    .card {
        background: #ffffff; border: 1px solid #dbe4f0;
        border-radius: 14px; padding: 1.2rem 1.4rem;
        margin-bottom: 1rem;
    }
    .card-top { border-left: 4px solid #7c9fd6; }

    /* ── Badges ── */
    .best-badge {
        background: #7c9fd6; color: #fff; font-size: 11px; font-weight: 600;
        padding: 3px 12px; border-radius: 20px; display: inline-block; margin-bottom: 8px;
    }
    .verdict { font-size: 12px; padding: 3px 12px; border-radius: 20px;
               display: inline-block; margin: 4px 0; font-weight: 500; }
    .v-strong { background: #e8f1ff; color: #355c9a; }
    .v-good   { background: #e7f7ef; color: #166534; }
    .v-partial{ background: #fff4db; color: #a16207; }
    .v-weak   { background: #ffe4e6; color: #be123c; }

    /* ── Tags ── */
    .tag { font-size: 11px; padding: 3px 10px; border-radius: 20px;
           margin: 2px; display: inline-block; }
    .tag-s { background: #e8f1ff; color: #355c9a; }
    .tag-g { background: #ffe4e6; color: #be123c; }

    /* ── Score ── */
    .score-big { font-size: 2.4rem; font-weight: 700; line-height: 1; text-align: center; }
    .score-label { font-size: 0.95rem; color: #64748b; text-align: center; }

    /* ── File pill ── */
    .file-pill {
        display: flex; align-items: center; gap: 10px;
        background: #f8fbff; border-radius: 8px;
        padding: 6px 12px; margin-bottom: 6px; font-size: 13px;
        border: 1px solid #dbe4f0;
    }
    .file-pill-icon { font-size: 16px; }
    .file-pill-name { flex: 1; font-weight: 500; }
    .file-pill-size { color: #64748b; font-size: 11px; }

    /* ── Status indicator ── */
    .status-dot {
        display: inline-block; width: 8px; height: 8px;
        border-radius: 50%; margin-right: 6px;
    }
    .dot-green { background: #4ade80; }
    .dot-red   { background: #f87171; }
    .dot-amber { background: #fbbf24; }

    /* ── Metric cards ── */
    [data-testid="stMetric"] {
        background: #ffffff; border: 1px solid #dbe4f0;
        border-radius: 10px; padding: 0.8rem 1rem;
    }

    /* ── Buttons ── */
    .stButton > button {
        border-radius: 10px; font-weight: 500;
        transition: all 0.15s;
    }
    .stButton > button[kind="primary"] {
        background: #b9d8ff !important; color: #1e3a8a !important;
        border: 1px solid #9ec5fe !important;
    }
    .stButton > button[kind="primary"]:hover {
        background: #9ec5fe !important;
    }

    /* ── Reset button danger style ── */
    .reset-btn > button {
        background: #ffe4e6 !important; color: #be123c !important;
        border: 1px solid #fda4af !important;
    }
    .reset-btn > button:hover {
        background: #fecdd3 !important;
    }

    /* ── Info box ── */
    .info-box {
        background: #e8f1ff; border-left: 4px solid #7c9fd6;
        border-radius: 0 8px 8px 0; padding: 10px 14px;
        font-size: 13px; color: #355c9a; margin-bottom: 12px;
    }
    .warn-box {
        background: #fff4db; border-left: 4px solid #d4a017;
        border-radius: 0 8px 8px 0; padding: 10px 14px;
        font-size: 13px; color: #8a5a00; margin-bottom: 12px;
    }
    .success-box {
        background: #e7f7ef; border-left: 4px solid #3ba776;
        border-radius: 0 8px 8px 0; padding: 10px 14px;
        font-size: 13px; color: #166534; margin-bottom: 12px;
    }
</style>
""", unsafe_allow_html=True)

# ─── CONSTANTS ────────────────────────────────────────────────────────────────
QWEN_MODEL          = "qwen/qwen3-32b"
FINAL_CHUNK_SIZE    = 700
FINAL_CHUNK_OVERLAP = 175
FINAL_K             = 6
REQUEST_DELAY       = 1

SYSTEM_PROMPT = """You are an HR screening assistant.
Your task is to compare a candidate's resume with a job description.

Focus on:
1. Relevant technical skills
2. Relevant work experience
3. Overall suitability for the role

Return the answer in a clear and concise format with:
- Candidate Summary
- Matching Skills
- Missing / Weak Areas
- Suitability Score out of 10
"""

# ─── HELPERS ──────────────────────────────────────────────────────────────────
def read_pdf_text(path: str) -> str:
    reader = PdfReader(str(path))
    return "\n".join(p.extract_text() or "" for p in reader.pages)

def read_txt_text(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def read_docx_text(path: str) -> str:
    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())

def read_uploaded_file(f) -> str:
    """Read a Streamlit UploadedFile to text, handling PDF, TXT, and DOCX."""
    suffix = Path(f.name).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(f.read())
        tmp_path = tmp.name
    try:
        if suffix == ".pdf":
            return read_pdf_text(tmp_path)
        if suffix == ".docx":
            return read_docx_text(tmp_path)
        return read_txt_text(tmp_path)
    finally:
        os.unlink(tmp_path)

def preprocess_text(text: str) -> str:
    if not isinstance(text, str):
        text = str(text)
    text = text.replace("\xa0"," ").replace("\n"," ").replace("\t"," ").lower()
    text = re.sub(r'(\+?\d[\d\-\(\) ]{7,}\d)', ' ', text)
    text = re.sub(r'[^a-z0-9/\+\#\-\.\s]', ' ', text)
    text = re.sub(r'[\.]{2,}', ' ', text)
    text = re.sub(r'[-]{2,}', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def extract_score(text: str):
    if not isinstance(text, str):
        return None
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL|re.IGNORECASE)
    cleaned = text.replace("**","").replace("__","").replace("*","").replace("\r","\n")
    patterns = [
        r"suitability\s*score\s*[:\-–]?\s*(\d+(?:\.\d+)?)\s*/\s*10",
        r"suitability\s*score\s*[:\-–]?\s*(\d+(?:\.\d+)?)\s*out\s*of\s*(?:10|ten)",
        r"suitability\s*score\s*[:\-–]?\s*(\d+(?:\.\d+)?)",
        r"(?:^|\n)\s*(?:overall\s+)?(?:final\s+)?(?:score|rating)\s*[:\-–]?\s*(\d+(?:\.\d+)?)\s*/\s*10",
        r"(?:^|\n)\s*(?:overall\s+)?(?:final\s+)?(?:score|rating)\s*[:\-–]?\s*(\d+(?:\.\d+)?)",
        r"[\(\[]\s*(\d+(?:\.\d+)?)\s*/\s*10\s*[\)\]]",
    ]
    for p in patterns:
        m = re.search(p, cleaned, re.IGNORECASE)
        if m:
            val = float(m.group(1))
            if 0 <= val <= 10:
                return val
    fb = re.findall(r"(?:suitability|score|rating)[^0-9\n]{0,20}(\d+(?:\.\d+)?)", cleaned, re.IGNORECASE)
    for m in reversed(fb):
        val = float(m)
        if 0 <= val <= 10:
            return val
    return None

def extract_section(text: str, names: list) -> str:
    if not isinstance(text, str):
        return ""
    cleaned = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL|re.IGNORECASE).strip()
    escaped = [re.escape(n) for n in names]
    pattern = (
        r"(?is)(?:^|\n)\s*(%s)\s*:\s*(.*?)"
        r"(?=\n\s*(?:Candidate Summary|Matching Skills|Missing\s*/\s*Weak Areas"
        r"|Missing Areas|Weak Areas|Suitability Score|Overall Score|Final Score)\s*:|\Z)"
    ) % "|".join(escaped)
    m = re.search(pattern, cleaned)
    if m:
        return re.sub(r"\s+", " ", m.group(2)).strip(" -\n\t")
    return ""

def build_user_prompt(jd: str, name: str, resume: str) -> str:
    return (f"Job Description:\n{jd}\n\nCandidate Resume File Name:\n{name}\n\n"
            f"Candidate Resume Text:\n{resume}\n\nPlease evaluate this candidate.\n\n"
            "Format your answer as:\nCandidate Summary:\nMatching Skills:\n"
            "Missing / Weak Areas:\nSuitability Score:")

def build_rag_prompt(jd: str, name: str, context: str) -> str:
    return (f"Job Description:\n{jd}\n\nCandidate Resume File Name:\n{name}\n\n"
            f"Most Relevant Resume Sections (retrieved via RAG):\n{context}\n\n"
            "Based on the retrieved resume sections above, please evaluate this candidate.\n\n"
            "Format your answer as:\nCandidate Summary:\nMatching Skills:\n"
            "Missing / Weak Areas:\nSuitability Score:")

def verdict_label(score):
    if score is None: return "Unknown", "v-weak"
    if score >= 8:    return "Strong Match", "v-strong"
    if score >= 6:    return "Good Match", "v-good"
    if score >= 4:    return "Partial Match", "v-partial"
    return "Weak Match", "v-weak"

def score_color(score):
    if score is None: return "#9a9895"
    if score >= 8:    return "#1c6641"
    if score >= 6:    return "#1e3f6e"
    if score >= 4:    return "#92400e"
    return "#991b1b"

def fmt_name(filename: str) -> str:
    """Strip extension for display."""
    return Path(filename).stem.replace("_", " ").replace("-", " ")

# ─── SESSION STATE ─────────────────────────────────────────────────────────────
_defaults = {
    "groq_key": "",
    "jd_files": {},        # name → preprocessed text
    "cv_files": {},        # name → preprocessed text
    "raw_jd_texts": {},    # name → raw text
    "raw_cv_texts": {},    # name → raw text
    "results": {},
    "embedding_model": None,
    "chat_history": [],
    "analysis_done": False,
    "files_loaded": False, # tracks whether Load & Preprocess has run
}
for k, v in _defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ─── RESET HELPER ─────────────────────────────────────────────────────────────
def reset_all():
    """
    Fully wipe all uploaded file data and results from session state.
    This is the CORRECT way to clear — pressing X on the file uploader
    only removes the widget value but does NOT clear session_state data.
    Always call this before loading a new batch of files.
    """
    for k in ["jd_files", "cv_files", "raw_jd_texts", "raw_cv_texts",
              "results", "chat_history"]:
        st.session_state[k] = {} if k != "chat_history" else []
    st.session_state["analysis_done"] = False
    st.session_state["files_loaded"]  = False
    st.rerun()

# ─── SIDEBAR ──────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## 🔍 TalentLens")
    st.caption("Qwen3-32B via Groq · FAISS RAG")
    st.divider()

    # ── API Key ──
    st.markdown("### 🔑 Groq API Key")
    key_val = st.text_input(
        "API Key", type="password",
        value=st.session_state.groq_key,
        label_visibility="collapsed",
        placeholder="Enter your Groq API key…"
    )
    col_save, col_status = st.columns([2, 1])
    with col_save:
        if st.button("💾 Save Key", use_container_width=True):
            st.session_state.groq_key = key_val.strip()
            st.success("Saved!")
    with col_status:
        if st.session_state.groq_key:
            st.markdown('<span class="status-dot dot-green"></span>OK', unsafe_allow_html=True)
        else:
            st.markdown('<span class="status-dot dot-red"></span>Not set', unsafe_allow_html=True)

    st.divider()

    # ── File Upload ──
    st.markdown("### 📁 Upload Files")

    # ⚠️  X BUTTON FIX — EXPLANATION:
    # Streamlit's file_uploader X button removes the file from the widget
    # but does NOT clear st.session_state. So stale data from previous runs
    # remains in memory and gets used in the next screening.
    #
    # The fix: we do NOT auto-process on upload. The user must explicitly
    # click "Load & Preprocess" to write into session_state. And we provide
    # a "Reset All Data" button that properly wipes session_state.
    # This guarantees: press X → click Reset → upload new files → Load → clean slate.

    jd_uploads = st.file_uploader(
        "Job Descriptions (PDF / TXT / DOCX)",
        type=["pdf", "txt", "docx"],
        accept_multiple_files=True,
        help="Upload one or more job description files."
    )
    cv_uploads = st.file_uploader(
        "Candidate CVs (PDF / TXT / DOCX)",
        type=["txt", "pdf", "docx"],
        accept_multiple_files=True,
        help="Upload all candidate CV files you want to screen."
    )

    st.divider()

    # ── Load & Preprocess ──
    if st.button("⚙️ Load & Preprocess", use_container_width=True, type="primary"):
        if not jd_uploads or not cv_uploads:
            st.warning("⚠️ Upload at least one JD and one CV first.")
        else:
            with st.spinner("Reading and preprocessing files…"):
                jd_files, raw_jd, cv_files, raw_cv = {}, {}, {}, {}
                errors = []

                for f in jd_uploads:
                    try:
                        raw = read_uploaded_file(f)
                        raw_jd[f.name]  = raw
                        jd_files[f.name] = preprocess_text(raw)
                    except Exception as e:
                        errors.append(f"JD '{f.name}': {e}")

                for f in cv_uploads:
                    try:
                        raw = read_uploaded_file(f)
                        raw_cv[f.name]  = raw
                        cv_files[f.name] = preprocess_text(raw)
                    except Exception as e:
                        errors.append(f"CV '{f.name}': {e}")

                st.session_state.update({
                    "jd_files":      jd_files,
                    "raw_jd_texts":  raw_jd,
                    "cv_files":      cv_files,
                    "raw_cv_texts":  raw_cv,
                    "results":       {},
                    "analysis_done": False,
                    "files_loaded":  True,
                })

            if errors:
                for e in errors:
                    st.error(e)
            else:
                st.success(f"✅ {len(jd_files)} JD(s) · {len(cv_files)} CV(s) loaded")

    # ── Status ──
    n_jd = len(st.session_state.jd_files)
    n_cv = len(st.session_state.cv_files)
    if st.session_state.files_loaded:
        st.markdown(
            f'<div style="font-size:12px;color:#9a9895;margin-top:4px">'
            f'<span class="status-dot dot-green"></span>'
            f'Loaded: {n_jd} JD · {n_cv} CVs</div>',
            unsafe_allow_html=True
        )
    else:
        st.markdown(
            '<div style="font-size:12px;color:#9a9895;margin-top:4px">'
            '<span class="status-dot dot-amber"></span>'
            'No files loaded yet</div>',
            unsafe_allow_html=True
        )

    st.divider()

    # ── ⚠️ RESET BUTTON — the real X-button fix ──
    st.markdown("### 🗑️ Reset")
    st.caption(
        "Pressing X on a file only removes it from the uploader widget — "
        "it does **not** clear the data already loaded into memory. "
        "Use this button to fully wipe everything before starting fresh."
    )
    with st.container():
        st.markdown('<div class="reset-btn">', unsafe_allow_html=True)
        if st.button("🗑️ Reset All Data", use_container_width=True):
            reset_all()
        st.markdown('</div>', unsafe_allow_html=True)

    st.divider()
    st.caption("TalentLens v3 · FAISS RAG · Qwen3-32B")

# ─── TABS ─────────────────────────────────────────────────────────────────────
tab_run, tab_res, tab_chat = st.tabs(["🚀 Run Screening", "📊 Results", "💬 Chat"])

# ══════════════════════════════════════════════════════════════════════════════
# TAB 1 — RUN SCREENING
# ══════════════════════════════════════════════════════════════════════════════
with tab_run:
    st.markdown("## CV Screening")

    if not st.session_state.groq_key:
        st.markdown(
            '<div class="warn-box">👈 <strong>Step 1:</strong> Add your Groq API key in the sidebar.</div>',
            unsafe_allow_html=True
        )
    elif not st.session_state.files_loaded:
        st.markdown(
            '<div class="info-box">👈 <strong>Step 2:</strong> Upload your JD and CV files in the sidebar, '
            'then click <strong>Load &amp; Preprocess</strong>.</div>',
            unsafe_allow_html=True
        )
    else:
        # ── File summary ──
        col_jd, col_cv = st.columns(2)
        with col_jd:
            st.markdown("**📄 Job Descriptions**")
            for n in st.session_state.jd_files:
                st.markdown(
                    f'<div class="file-pill">'
                    f'<span class="file-pill-icon">📄</span>'
                    f'<span class="file-pill-name">{fmt_name(n)}</span>'
                    f'<span class="file-pill-size">{len(st.session_state.raw_jd_texts[n]):,} chars</span>'
                    f'</div>',
                    unsafe_allow_html=True
                )
        with col_cv:
            st.markdown("**👤 Candidate CVs**")
            for n in st.session_state.cv_files:
                st.markdown(
                    f'<div class="file-pill">'
                    f'<span class="file-pill-icon">👤</span>'
                    f'<span class="file-pill-name">{fmt_name(n)}</span>'
                    f'<span class="file-pill-size">{len(st.session_state.raw_cv_texts[n]):,} chars</span>'
                    f'</div>',
                    unsafe_allow_html=True
                )

        st.divider()

        # ── Options ──
        col_opt1, col_opt2 = st.columns([2, 1])
        with col_opt1:
            use_rag = st.checkbox(
                "✅ Use RAG (recommended)",
                value=True,
                help="Uses FAISS vector store to retrieve the most relevant sections "
                     "of each CV before scoring. Produces better, more focused results."
            )
        with col_opt2:
            delay = st.number_input(
                "Delay between calls (s)",
                min_value=0, max_value=5, value=REQUEST_DELAY, step=1,
                help="Rate-limit delay between API calls. Increase if you hit 429 errors."
            )

        # ── Run button ──
        total_pairs = len(st.session_state.jd_files) * len(st.session_state.cv_files)
        st.markdown(
            f'<div class="info-box">Ready to screen '
            f'<strong>{len(st.session_state.cv_files)} candidate(s)</strong> against '
            f'<strong>{len(st.session_state.jd_files)} job description(s)</strong> '
            f'— {total_pairs} API call(s) total.</div>',
            unsafe_allow_html=True
        )

        if st.button("▶️ Run Screening", type="primary", use_container_width=True):
            groq_client = Groq(api_key=st.session_state.groq_key)

            # Load embedding model once and cache in session state
            with st.spinner("Loading embedding model (first run may take ~30s)…"):
                if st.session_state.embedding_model is None:
                    st.session_state.embedding_model = HuggingFaceEmbeddings(
                        model_name="sentence-transformers/all-MiniLM-L6-v2"
                    )
            emb = st.session_state.embedding_model

            # Build FAISS chunks once for all CVs
            with st.spinner("Building FAISS vector index…"):
                docs = [
                    Document(page_content=t, metadata={"source": n})
                    for n, t in st.session_state.cv_files.items()
                ]
                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=FINAL_CHUNK_SIZE,
                    chunk_overlap=FINAL_CHUNK_OVERLAP
                )
                chunks = splitter.split_documents(docs)

            def get_rag_context(cv_name: str, query: str) -> str:
                cand_chunks = [c for c in chunks if c.metadata["source"] == cv_name]
                if not cand_chunks:
                    return ""
                vs  = FAISS.from_documents(cand_chunks, emb)
                ret = vs.as_retriever(search_kwargs={"k": min(FINAL_K, len(cand_chunks))})
                return "\n\n".join(d.page_content for d in ret.invoke(query))

            def call_llm(prompt: str) -> str:
                r = groq_client.chat.completions.create(
                    model=QWEN_MODEL,
                    messages=[
                        {"role": "system", "content": SYSTEM_PROMPT},
                        {"role": "user",   "content": prompt}
                    ],
                    temperature=0.1,
                    max_tokens=1024,
                )
                return r.choices[0].message.content

            prog     = st.progress(0, text="Starting…")
            step     = 0
            all_results = {}

            for jd_name, cleaned_jd in st.session_state.jd_files.items():
                jd_res = []
                for cv_name, cleaned_cv in st.session_state.cv_files.items():
                    step += 1
                    prog.progress(
                        step / total_pairs,
                        text=f"({step}/{total_pairs}) Scoring {fmt_name(cv_name)}…"
                    )

                    if use_rag:
                        ctx    = get_rag_context(cv_name, cleaned_jd[:500])
                        prompt = build_rag_prompt(cleaned_jd, cv_name, ctx)
                    else:
                        prompt = build_user_prompt(cleaned_jd, cv_name, cleaned_cv)

                    try:
                        output = call_llm(prompt)
                    except Exception as e:
                        output = f"Error: {e}"

                    jd_res.append({
                        "cv_name": cv_name,
                        "score":   extract_score(output),
                        "output":  output,
                        "summary": extract_section(output, ["Candidate Summary"]),
                        "skills":  extract_section(output, ["Matching Skills"]),
                        "gaps":    extract_section(output, ["Missing / Weak Areas", "Missing Areas", "Weak Areas"]),
                        "mode":    "RAG" if use_rag else "Direct",
                    })
                    time.sleep(delay)

                jd_res.sort(
                    key=lambda x: x["score"] if x["score"] is not None else -1,
                    reverse=True
                )
                all_results[jd_name] = jd_res

            prog.empty()
            st.session_state.results      = all_results
            st.session_state.analysis_done = True

            # Quick winner summary
            first_jd = list(all_results.keys())[0]
            best     = all_results[first_jd][0]
            st.markdown(
                f'<div class="success-box">✅ Screening complete! '
                f'Top candidate: <strong>{fmt_name(best["cv_name"])}</strong> '
                f'— Score: <strong>{best["score"]}/10</strong>. '
                f'Open the <strong>Results</strong> tab to see full rankings.</div>',
                unsafe_allow_html=True
            )

# ══════════════════════════════════════════════════════════════════════════════
# TAB 2 — RESULTS
# ══════════════════════════════════════════════════════════════════════════════
with tab_res:
    st.markdown("## Results")

    if not st.session_state.analysis_done:
        st.markdown(
            '<div class="warn-box">⚠️ No results yet — run screening in the '
            '<strong>Run Screening</strong> tab first.</div>',
            unsafe_allow_html=True
        )
    else:
        res_dict = st.session_state.results
        jd_names = list(res_dict.keys())

        # JD selector if multiple JDs were uploaded
        if len(jd_names) > 1:
            sel_jd = st.selectbox(
                "Select Job Description",
                jd_names,
                format_func=fmt_name
            )
        else:
            sel_jd = jd_names[0]

        jd_res = res_dict[sel_jd]
        scores = [r["score"] for r in jd_res if r["score"] is not None]

        # ── Top candidate banner ──
        if jd_res:
            best = jd_res[0]
            verdict, vcls = verdict_label(best["score"])
            color = score_color(best["score"])
            st.markdown(
                f'<div class="card card-top">'
                f'<span class="best-badge">⭐ Best Match</span><br>'
                f'<span style="font-size:1.3rem;font-weight:700">{fmt_name(best["cv_name"])}</span>'
                f'<span style="color:#9a9895;margin-left:10px">Score: '
                f'<span style="color:{color};font-weight:700">'
                f'{best["score"] if best["score"] is not None else "N/A"}/10</span></span><br>'
                f'<span class="verdict {vcls}">{verdict}</span>'
                f'{"<br><small>" + best["summary"] + "</small>" if best["summary"] else ""}'
                f'</div>',
                unsafe_allow_html=True
            )

        # ── Metrics row ──
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("CVs Screened",  len(jd_res))
        m2.metric("Best Score",    f"{max(scores):.1f}/10"                      if scores else "—")
        m3.metric("Avg Score",     f"{sum(scores)/len(scores):.1f}/10"          if scores else "—")
        m4.metric("Top Candidate", fmt_name(jd_res[0]["cv_name"])               if jd_res else "—")

        st.divider()

        # ── Bar chart ──
        st.markdown("**Score Comparison**")
        chart_data = pd.DataFrame({
            "Candidate": [fmt_name(r["cv_name"]) for r in jd_res],
            "Score /10": [r["score"] if r["score"] is not None else 0 for r in jd_res],
        }).set_index("Candidate")
        st.bar_chart(chart_data, color="#1c6641")

        st.divider()
        st.markdown("### Candidate Cards")

        # ── Candidate expander cards ──
        for i, r in enumerate(jd_res):
            verdict, vcls = verdict_label(r["score"])
            color         = score_color(r["score"])
            rank_icon     = "🥇" if i == 0 else ("🥈" if i == 1 else ("🥉" if i == 2 else f"#{i+1}"))
            score_str     = f"{r['score']}/10" if r["score"] is not None else "N/A"
            mode_badge    = f' <span style="font-size:11px;background:#f2f1ed;padding:2px 8px;border-radius:10px;color:#555">{r.get("mode","")}</span>'

            with st.expander(
                f"{rank_icon}  {fmt_name(r['cv_name'])}  —  {score_str}",
                expanded=(i == 0)
            ):
                left, right = st.columns([3, 1])

                with left:
                    if i == 0:
                        st.markdown('<span class="best-badge">⭐ Best Match</span>', unsafe_allow_html=True)
                    st.markdown(
                        f'<span class="verdict {vcls}">{verdict}</span>{mode_badge}',
                        unsafe_allow_html=True
                    )

                    if r["summary"]:
                        st.markdown(f"**Summary:** {r['summary']}")

                    if r["skills"]:
                        tags = [s.strip() for s in re.split(r"[,;•\-\n]+", r["skills"]) if s.strip()][:8]
                        st.markdown("**Matching Skills:**")
                        st.markdown(
                            " ".join(f'<span class="tag tag-s">{t}</span>' for t in tags),
                            unsafe_allow_html=True
                        )

                    if r["gaps"]:
                        gtags = [g.strip() for g in re.split(r"[,;•\-\n]+", r["gaps"]) if g.strip()][:5]
                        st.markdown("**Gaps / Weak Areas:**")
                        st.markdown(
                            " ".join(f'<span class="tag tag-g">{g}</span>' for g in gtags),
                            unsafe_allow_html=True
                        )

                with right:
                    st.markdown(
                        f'<div class="score-big" style="color:{color}">'
                        f'{r["score"] if r["score"] is not None else "—"}'
                        f'</div>'
                        f'<div class="score-label">/10</div>',
                        unsafe_allow_html=True
                    )
                    if r["score"] is not None:
                        st.progress(float(r["score"]) / 10.0)

                with st.expander("📋 Full AI Evaluation"):
                    clean = re.sub(
                        r"<think>.*?</think>", "", r["output"],
                        flags=re.DOTALL|re.IGNORECASE
                    ).strip()
                    st.markdown(clean)

        st.divider()

        # ── Download ──
        rows = [{
            "Rank":        i + 1,
            "Candidate":   fmt_name(r["cv_name"]),
            "File":        r["cv_name"],
            "Score /10":   r["score"],
            "Verdict":     verdict_label(r["score"])[0],
            "Mode":        r.get("mode", ""),
            "Summary":     (r["summary"] or "")[:200],
            "Skills":      (r["skills"]  or "")[:200],
            "Gaps":        (r["gaps"]    or "")[:200],
        } for i, r in enumerate(jd_res)]

        st.download_button(
            "⬇️ Download Full Results CSV",
            pd.DataFrame(rows).to_csv(index=False),
            file_name=f"talentlens_{Path(sel_jd).stem}.csv",
            mime="text/csv",
            use_container_width=True,
        )

# ══════════════════════════════════════════════════════════════════════════════
# TAB 3 — CHAT
# ══════════════════════════════════════════════════════════════════════════════
with tab_chat:
    st.markdown("## 💬 Ask About Candidates")
    st.caption("Ask anything — CV details, compare candidates, generate interview questions, etc.")

    if not st.session_state.groq_key:
        st.markdown(
            '<div class="warn-box">👈 Set your Groq API key in the sidebar first.</div>',
            unsafe_allow_html=True
        )
    else:
        # ── Build context — COMPACT to avoid 413 token errors ──
        # Raw CV/JD text is NOT included — only screening result summaries.
        # This keeps the system prompt well under 3000 tokens even with 4 JDs × 7 CVs.
        def build_chat_context() -> str:
            ctx = ""
            # JD names only (not full text) — saves ~2000 tokens per JD
            if st.session_state.raw_jd_texts:
                ctx += "=== JOB DESCRIPTIONS LOADED ===\n"
                for n in st.session_state.raw_jd_texts:
                    ctx += f"  - {n}\n"
            # CV names only
            if st.session_state.raw_cv_texts:
                ctx += "\n=== CANDIDATE CVs LOADED ===\n"
                for n in st.session_state.raw_cv_texts:
                    ctx += f"  - {n}\n"
            # Full screening results — scores, verdicts, summaries, skills, gaps
            if st.session_state.results:
                ctx += "\n=== SCREENING RESULTS ===\n"
                for jd_n, res in st.session_state.results.items():
                    ctx += f"\nJob Description: {jd_n}\n"
                    for r in res:
                        verdict, _ = verdict_label(r["score"])
                        ctx += (
                            f"  Candidate: {fmt_name(r['cv_name'])} | "
                            f"Score: {r['score']}/10 | Verdict: {verdict}\n"
                        )
                        if r.get("summary"):
                            ctx += f"    Summary: {(r['summary'] or '')[:200]}\n"
                        if r.get("skills"):
                            ctx += f"    Skills: {(r['skills'] or '')[:150]}\n"
                        if r.get("gaps"):
                            ctx += f"    Gaps: {(r['gaps'] or '')[:150]}\n"
            return ctx.strip()

        # ── Central function that calls the LLM and saves reply ──
        def run_chat_llm(question: str):
            """Send question + history to Groq, append reply to chat history."""
            groq_client = Groq(api_key=st.session_state.groq_key)
            sys_content = (
                "You are an expert HR recruitment assistant. "
                "Use the screening data below to answer questions accurately. "
                "Cite candidate names and scores where relevant. Be concise.\n\n"
                + build_chat_context()
            )
            msgs = [{"role": "system", "content": sys_content}]
            # Only keep last 6 chat turns to avoid history bloat
            recent_history = st.session_state.chat_history[-6:]
            msgs += [{"role": m["role"], "content": m["content"]} for m in recent_history]
            try:
                resp = groq_client.chat.completions.create(
                    model=QWEN_MODEL,
                    messages=msgs,
                    temperature=0.3,
                    max_tokens=800,  # reduced from 1500 — keeps response under TPM limit
                )
                reply = resp.choices[0].message.content
                reply = re.sub(
                    r"<think>.*?</think>", "", reply,
                    flags=re.DOTALL | re.IGNORECASE
                ).strip()
            except Exception as e:
                err = str(e)
                if "413" in err or "rate_limit" in err.lower() or "tokens" in err.lower():
                    reply = (
                        "⚠️ The request was too large for the free API tier. "
                        "Try asking a more specific question, or clear the chat history to free up space."
                    )
                else:
                    reply = f"❌ Error: {e}"
            st.session_state.chat_history.append({"role": "assistant", "content": reply})

        # ── Quick prompt buttons — only shown when chat is empty ──
        if not st.session_state.chat_history:
            if not st.session_state.files_loaded:
                st.markdown(
                    '<div class="info-box">ℹ️ Upload files and run screening first '
                    'to get context-aware answers. You can still chat without files.</div>',
                    unsafe_allow_html=True
                )

            st.markdown("**💡 Try asking:**")
            quick_prompts = [
                "Who is the best candidate and why?",
                "Compare the top 2 candidates side by side.",
                "Generate 5 interview questions for the top candidate.",
                "What skills are missing across all candidates?",
                "Which candidate would need the least training?",
                "Summarise all candidates in a table.",
            ]
            cols = st.columns(3)
            for idx, q in enumerate(quick_prompts):
                with cols[idx % 3]:
                    if st.button(q, use_container_width=True, key=f"qp_{idx}"):
                        # FIX: append question AND immediately get the answer
                        # before rerunning — otherwise the page reruns with no reply
                        st.session_state.chat_history.append({"role": "user", "content": q})
                        run_chat_llm(q)
                        st.rerun()

        # ── Render chat history ──
        for msg in st.session_state.chat_history:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        # ── Chat input (typed messages) ──
        user_msg = st.chat_input("Ask about any candidate, skill gap, or comparison…")
        if user_msg:
            st.session_state.chat_history.append({"role": "user", "content": user_msg})
            with st.chat_message("assistant"):
                with st.spinner("Thinking…"):
                    run_chat_llm(user_msg)
            # Show the reply that was just appended
            with st.chat_message("assistant"):
                st.markdown(st.session_state.chat_history[-1]["content"])
            st.rerun()

        # ── Clear chat ──
        if st.session_state.chat_history:
            st.divider()
            if st.button("🗑️ Clear Chat History", use_container_width=False):
                st.session_state.chat_history = []
                st.rerun()
