import os
import json
import re
import time
import tempfile
from pathlib import Path
from typing import Dict, List, Any

from pypdf import PdfReader
from docx import Document as DocxDocument
from groq import Groq
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings

VERSION = "v3.7-hr-shared-results"  # bump this when you redeploy to confirm Render picked up the new file

QWEN_MODEL = "qwen/qwen3-32b"
FINAL_CHUNK_SIZE = 700
FINAL_CHUNK_OVERLAP = 175
FINAL_K = 6
SYSTEM_PROMPT = """You are a senior HR screening assistant with expertise in 
technical hiring for semiconductor and hardware engineering roles.

You will be given a job description (JD) and a candidate's resume. Your task 
is to produce a structured screening report.

Follow these rules strictly:
- Only assess skills and experience explicitly stated in the resume. Do not infer 
  or assume unstated competencies.
- Distinguish between hands-on experience and exposure/familiarity.
- Weight criteria in this order: (1) Must-have technical skills, (2) Years of 
  relevant experience, (3) Nice-to-have skills, (4) Education.
- If a requirement from the JD has no match in the resume, explicitly flag it 
  as a gap — do not omit it.

Return your report in this exact structure:

**Candidate Summary** (2-3 sentences: role fit at a glance)

**Must-Have Skills Match**
| Requirement | Found in Resume? | Evidence |
|---|---|---|

**Nice-to-Have Skills Match**
| Requirement | Found? | Evidence |
|---|---|---|

**Experience Assessment**
- Years of relevant experience: 
- Seniority alignment:
- Notable projects/achievements:

**Key Gaps**
(List missing or weak areas with specific JD requirements they fail to meet)

**Suitability Score: X/10**
Scoring rationale: (2-3 sentences explaining the score)

**Recommendation:** Advance / Hold / Reject
"""

_CHAT_SYSTEM_PROMPT = """You are an AI assistant embedded in a CV screening 
tool used by an HR team to evaluate engineering candidates.

You have access to:
- The job description
- Candidate resumes (raw text)
- Pre-generated screening reports for each candidate

Your behaviour rules:
- Answer only from the provided context. Never fabricate experience or skills 
  a candidate does not have.
- When comparing candidates, use specific evidence (e.g. "Candidate A has 5 
  years of UVM experience vs Candidate B's 2 years").
- If asked to rank candidates, justify each placement with data from their reports.
- If the context is insufficient to answer confidently, say exactly what 
  information is missing rather than guessing.
- Keep answers concise but evidence-based. Avoid vague qualifiers like 
  "seems strong" — prefer "has 3 years of hands-on SystemVerilog per resume".
"""

_EMBEDDINGS = None


def get_embedding_model():
    global _EMBEDDINGS
    if _EMBEDDINGS is None:
        _EMBEDDINGS = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    return _EMBEDDINGS


def read_pdf_text(path: str) -> str:
    """Extract text from PDF preserving visual reading order.

    Uses layout mode which respects column structure, preventing words from
    different columns/sections being incorrectly concatenated by pypdf.
    Falls back to plain mode if layout extraction fails.
    """
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        try:
            text = page.extract_text(extraction_mode="layout") or ""
        except Exception:
            text = page.extract_text() or ""  # plain mode fallback
        pages.append(text)
    return "\n".join(pages)


def read_txt_text(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def read_docx_text(path: str) -> str:
    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())


def read_uploaded_file(uploaded_file) -> str:
    suffix = Path(uploaded_file.name).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.getbuffer())
        tmp_path = tmp.name
    try:
        if suffix == ".pdf":
            return read_pdf_text(tmp_path)
        if suffix == ".docx":
            return read_docx_text(tmp_path)
        if suffix == ".txt":
            return read_txt_text(tmp_path)
        raise ValueError(f"Unsupported file type: {suffix}")
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def preprocess_text(text: str) -> str:
    if not isinstance(text, str):
        text = str(text)
    text = text.replace("\xa0", " ").replace("\n", " ").replace("\t", " ").lower()
    text = re.sub(r'(\+?\d[\d\-\(\) ]{7,}\d)', ' ', text)
    text = re.sub(r'[^a-z0-9/\+\#\-\.\s]', ' ', text)
    text = re.sub(r'[\.]{2,}', ' ', text)
    text = re.sub(r'[-]{2,}', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def extract_score(text: str):
    if not isinstance(text, str):
        return None
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL | re.IGNORECASE)
    cleaned = text.replace("**", "").replace("__", "").replace("*", "").replace("\r", "\n")
    patterns = [
        r"suitability\s*score\s*[:\-–]?\s*(\d+(?:\.\d+)?)\s*/\s*10",
        r"suitability\s*score\s*[:\-–]?\s*(\d+(?:\.\d+)?)\s*out\s*of\s*(?:10|ten)",
        r"suitability\s*score\s*[:\-–]?\s*(\d+(?:\.\d+)?)",
        r"(?:^|\n)\s*(?:overall\s+)?(?:final\s+)?(?:score|rating)\s*[:\-–]?\s*(\d+(?:\.\d+)?)\s*/\s*10",
        r"(?:^|\n)\s*(?:overall\s+)?(?:final\s+)?(?:score|rating)\s*[:\-–]?\s*(\d+(?:\.\d+)?)\s*out\s*of\s*(?:10|ten)",
        r"(?:^|\n)\s*(?:overall\s+)?(?:final\s+)?(?:score|rating)\s*[:\-–]?\s*(\d+(?:\.\d+)?)",
        r"[\(\[]\s*(\d+(?:\.\d+)?)\s*/\s*10\s*[\)\]]",
    ]
    for pattern in patterns:
        m = re.search(pattern, cleaned, re.IGNORECASE)
        if m:
            val = float(m.group(1))
            if 0 <= val <= 10:
                return val
    fallback = re.findall(r"(?:suitability|score|rating)[^0-9\n]{0,20}(\d+(?:\.\d+)?)", cleaned, re.IGNORECASE)
    for item in reversed(fallback):
        val = float(item)
        if 0 <= val <= 10:
            return val
    return None


def extract_section(text: str, names: List[str]) -> str:
    if not isinstance(text, str):
        return ""
    cleaned = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL | re.IGNORECASE).strip()
    escaped = [re.escape(n) for n in names]
    pattern = (
        r"(?is)(?:^|\n)\s*(%s)\s*:\s*(.*?)"
        r"(?=\n\s*(?:Candidate Summary|Matching Skills|Missing\s*/\s*Weak Areas"
        r"|Missing Areas|Weak Areas|Suitability Score|Overall Score|Final Score)\s*:|\Z)"
    ) % "|".join(escaped)
    m = re.search(pattern, cleaned)
    if not m:
        return ""
    return re.sub(r"\s+", " ", m.group(2)).strip(" -\n\t")


def fmt_name(filename: str) -> str:
    return Path(filename).stem.replace("_", " ").replace("-", " ")


def verdict_label(score):
    if score is None:
        return "Unknown", "v-weak"
    if score >= 8:
        return "Strong Match", "v-strong"
    if score >= 6:
        return "Good Match", "v-good"
    if score >= 4:
        return "Partial Match", "v-partial"
    return "Weak Match", "v-weak"


def score_color(score):
    if score is None:
        return "#94a3b8"
    if score >= 8:
        return "#0f766e"
    if score >= 6:
        return "#1d4ed8"
    if score >= 4:
        return "#b45309"
    return "#b91c1c"


def build_user_prompt(jd: str, name: str, resume: str) -> str:
    return (
        f"Job Description:\n{jd}\n\nCandidate Resume File Name:\n{name}\n\n"
        f"Candidate Resume Text:\n{resume}\n\nPlease evaluate this candidate.\n\n"
        "Format your answer as:\nCandidate Summary:\nMatching Skills:\n"
        "Missing / Weak Areas:\nSuitability Score:"
    )


def build_rag_prompt(jd: str, name: str, context: str) -> str:
    return (
        f"Job Description:\n{jd}\n\nCandidate Resume File Name:\n{name}\n\n"
        f"Most Relevant Resume Sections (retrieved via RAG):\n{context}\n\n"
        "Based on the retrieved resume sections above, please evaluate this candidate.\n\n"
        "Format your answer as:\nCandidate Summary:\nMatching Skills:\n"
        "Missing / Weak Areas:\nSuitability Score:"
    )


def load_and_preprocess_files(jd_uploads, cv_uploads):
    jd_files, raw_jd, cv_files, raw_cv = {}, {}, {}, {}
    errors = []

    for f in jd_uploads or []:
        try:
            raw = read_uploaded_file(f)
            raw_jd[f.name] = raw
            jd_files[f.name] = preprocess_text(raw)
        except Exception as e:
            errors.append(f"JD '{f.name}': {e}")

    for f in cv_uploads or []:
        try:
            raw = read_uploaded_file(f)
            raw_cv[f.name] = raw
            cv_files[f.name] = preprocess_text(raw)
        except Exception as e:
            errors.append(f"CV '{f.name}': {e}")

    return {
        "jd_files": jd_files,
        "raw_jd_texts": raw_jd,
        "cv_files": cv_files,
        "raw_cv_texts": raw_cv,
        "errors": errors,
    }


def _build_chunks(cv_files: Dict[str, str]):
    documents = [Document(page_content=text, metadata={"source": name}) for name, text in cv_files.items()]
    splitter = RecursiveCharacterTextSplitter(chunk_size=FINAL_CHUNK_SIZE, chunk_overlap=FINAL_CHUNK_OVERLAP)
    return splitter.split_documents(documents)


def _get_rag_context(chunks, emb, cv_name: str, query: str) -> str:
    cand_chunks = [c for c in chunks if c.metadata.get("source") == cv_name]
    if not cand_chunks:
        return ""
    vectorstore = FAISS.from_documents(cand_chunks, emb)
    retriever = vectorstore.as_retriever(search_kwargs={"k": min(FINAL_K, len(cand_chunks))})
    return "\n\n".join(d.page_content for d in retriever.invoke(query))


def _call_llm(client: Groq, prompt: str, retries: int = 3, base_delay: float = 5.0) -> str:
    """Call the LLM with exponential backoff retry on rate-limit or server errors."""
    last_err = None
    for attempt in range(retries):
        try:
            response = client.chat.completions.create(
                model=QWEN_MODEL,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.1,
                max_tokens=1024,
            )
            return response.choices[0].message.content
        except Exception as e:
            last_err = e
            err_str = str(e).lower()
            # Only retry on rate limits or server errors, not auth/bad request
            if any(code in err_str for code in ["429", "500", "503", "rate_limit", "timeout"]):
                wait = base_delay * (2 ** attempt)  # 5s, 10s, 20s
                time.sleep(wait)
            else:
                raise  # non-retryable error, fail immediately
    raise RuntimeError(f"LLM call failed after {retries} retries: {last_err}")


# ── ANONYMISATION ─────────────────────────────────────────────────────────────
# NOTE: anonymise_text() must be called on RAW text (before preprocess_text),
# because preprocess_text lowercases and strips punctuation, which breaks
# capitalised-name detection and email/URL matching.
#
# v3.1 fix: do NOT globally redact every Title Case / ALL CAPS phrase.
# That was over-censoring useful CV content such as "DIPLOMA IN AI",
# "Python Programming", "National Day", "Basic SQL", and job titles.
# Names are now removed only from likely name positions such as CV headers,
# labelled "Name:" fields, emergency contacts, and title-prefixed names.

# ── Emails: standard + obfuscated variants ────────────────────────────────────
_EMAIL = re.compile(
    r"[\w.\-+]+\s*(?:@|\[at\]|\(at\)|＠)\s*[\w.\-]+\s*(?:\."
    r"|\[dot\]|\(dot\))\s*[a-z]{2,}",
    re.IGNORECASE,
)

# ── Phone numbers ─────────────────────────────────────────────────────────────
# Carefully avoids matching year ranges like 2021-2023.
# Dashed numbers only matched when first block starts with SG prefix (6/8/9).
_PHONE = re.compile(
    r"(?<!\d)(?:"
    r"\+\d{1,3}[\s\-.]?\d{4,5}[\s\-.]?\d{4,5}"   # +65 9123 4567
    r"|\(\d{1,4}\)[\s\-.]?\d{3,5}[\s\-.]?\d{3,5}" # (65) 9123 4567
    r"|[689]\d{7}"                                  # bare SG 8-digit (no separator)
    r"|[689]\d{3}[\s\.\-]\d{4}"                    # SG with separator: 9123-4567
    r"|\d{4,5}[\s\.]\d{4,5}"                       # other spaced/dotted (not dashed, avoids years)
    r")(?!\d)",
    re.IGNORECASE,
)

# ── Singapore postal code: 6-digit, optionally preceded by "Singapore" ────────
_SG_POSTAL = re.compile(r"\b(?:singapore\s+|s\()?\d{6}\b", re.IGNORECASE)

# ── Singapore block/unit address ──────────────────────────────────────────────
_SG_ADDR = re.compile(
    r"\b(?:blk|block)\s+\d+\b[^,\n]*"  # Blk 85 Street Name
    r"|\#\d+[\-\u2013]\d+",            # #10-01
    re.IGNORECASE,
)

# ── Generic street addresses ──────────────────────────────────────────────────
_STREET_ADDR = re.compile(
    r"\b\d+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,5}\s+"
    r"(?:Street|St|Road|Rd|Avenue|Ave|Drive|Dr|Lane|Ln|Way|Close|Crescent|Cres|"
    r"Place|Pl|Court|Ct|Boulevard|Blvd|Walk|Rise|View|Hill|Gardens|Park)\b"
    r"(?:\s+\d+)?",
    re.IGNORECASE,
)

# ── URLs: with or without http/https ──────────────────────────────────────────
_URL = re.compile(
    r"(?:https?://|www\.)\S+"
    r"|(?:linkedin\.com|github\.com|gitlab\.com|behance\.net|"
    r"dribbble\.com|portfolio\.|medium\.com/@?)\S*",
    re.IGNORECASE,
)

# ── Labelled names and title-prefixed names ───────────────────────────────────
_LABELLED_NAME = re.compile(
    r"(?im)^(\s*(?:[▪•\-]\s*)?(?:name|candidate\s+name|full\s+name)\s*[:\-]\s*)[^\n\r]+"
)
_NAME_WITH_TITLE = re.compile(
    r"\b(?:mr|mrs|ms|miss|dr|prof|sir|assoc\.?\s+prof)\.?\s+"
    r"[A-Z][A-Za-z'\-]{1,25}(?:\s+[A-Z][A-Za-z'\-]{1,25}){0,3}\b"
)

# ── Section headings / words that must never be treated as names ──────────────
_HEADER_STOPWORDS = {
    "contact", "coordonnées", "top skills", "principales compétences", "skills",
    "languages", "language", "experience", "professional experience", "education",
    "summary", "objective", "overview", "technical", "soft skill", "honors",
    "honors-awards", "honors & awards", "awards", "awards / achievements",
    "work experience", "relevant competencies / skills", "computer literacy",
    "communication skills", "management and leadership skills", "licensures & certifications",
    "emergency contacts", "official closed and non-sensitive", "official closed and non sensitive",
}

_NON_NAME_WORDS = {
    # CV / section words
    "official", "closed", "non", "sensitive", "contact", "education", "experience",
    "skills", "languages", "summary", "objective", "overview", "technical", "project",
    "projects", "achievement", "achievements", "awards", "honors", "activities",
    "volunteer", "certifications", "licensures", "competencies", "literacy",
    "communication", "management", "leadership", "relationship", "mother", "father",
    # tech / education words
    "diploma", "degree", "bachelor", "master", "engineering", "engineer", "science",
    "computer", "data", "ai", "machine", "learning", "python", "programming", "basic",
    "sql", "rag", "langchain", "nlp", "llm", "electronics", "electrical", "digital",
    "analog", "verilog", "systemverilog", "uvm", "rtl", "fpga", "soc", "asic",
    "vlsi", "eda", "linux", "perl", "tcl", "bash", "cadence", "synopsys",
    "mentor", "vcs", "verdi", "innovus", "primetime", "synthesis", "floorplan",
    "floorplanning", "place", "route", "timing", "power", "coverage", "simulation",
    # company / location / date words should not trigger name redaction alone
    "server", "staff", "part", "time", "national", "day", "defense", "audio", "setup",
}

# ── Institutions (specific enough to avoid eating education titles) ───────────
_INSTITUTION = re.compile(
    r"\b[A-Z][A-Za-z&.'\-]*(?:\s+[A-Z][A-Za-z&.'\-]*){0,5}\s+"
    r"(?:University|Polytechnic|College|Academy|School|Institute(?:\s+of\s+Technology)?)\b"
    r"(?:\s+of\s+[A-Z][A-Za-z&.'\-]*(?:\s+[A-Z][A-Za-z&.'\-]*){0,4})?"
    r"(?:,?\s+(?:Singapore|Malaysia|Vietnam|India|China|Korea|Japan|USA|United\s+States))?",
    re.IGNORECASE,
)

# ── Nationality / race / language markers ────────────────────────────────────
# Keep English because many JDs require English communication skills.
# Language words are replaced with [language], not [nationality], so the preview
# is easier to explain during testing.
_LANGUAGE = re.compile(
    r"\b(chinese|korean|japanese|mandarin|tamil|hokkien|cantonese|teochew|"
    r"hakka|marathi|hindi|malay)\b",
    re.IGNORECASE,
)
_NATIONALITY = re.compile(
    r"\b(singaporean|malaysian|indonesian|filipino|vietnamese|burmese|"
    r"thai|indian|eurasian|caucasian|american|british|australian|canadian)\b",
    re.IGNORECASE,
)

# ── Date of birth / age / marital status / national ID ───────────────────────
_NRIC = re.compile(r"\b[STFG]\d{7}[A-Z]\b", re.IGNORECASE)
_DOB = re.compile(
    r"\b(?:date\s+of\s+birth|d\.?o\.?b\.?|age|born\s+(?:in|on))"
    r"\s*[:\-]?\s*[\d/\-\w,\s]{0,30}",
    re.IGNORECASE,
)
_MARITAL = re.compile(
    r"\b(single|married|divorced|widowed|marital\s+status\s*[:\-]?\s*\w+)\b",
    re.IGNORECASE,
)
_RELATIONSHIP = re.compile(r"(?im)^(\s*(?:[▪•\-]\s*)?relationship\s*[:\-]\s*)[^\n\r]+")
_CONTACT_NO = re.compile(r"(?im)^(\s*(?:[▪•\-]\s*)?(?:contact\s+no|mobile|phone|home)\s*[:\-]\s*)[^\n\r]+")

# Remove repeated school document classification labels such as
# "Official (Closed) and Non-Sensitive". These are not useful to the LLM.
_DOCUMENT_LABEL = re.compile(
    r"(?im)^\s*official\s*\(?\s*closed\s*\)?\s*(?:and|&)\s*non[-\s]*sensitive\s*$"
)

# After phone/email/address replacements, collapse full contact/address lines so
# the preview does not become messy text like "mobile phone home phone email email".
_CONTACT_DETAIL_LINE = re.compile(
    r"(?im)^.*\b(?:mobile|phone|home|email|contact\s+no)\b.*(?:\[phone\]|\[email\]).*$"
)
_ADDRESS_DETAIL_LINE = re.compile(
    r"(?im)^.*(?:\[address\]|\[postal\]).*$"
)

_PRONOUN_MAP = {
    "he": "they", "she": "they",
    "him": "them", "her": "them",
    "his": "their", "hers": "theirs",
    "himself": "themself", "herself": "themself",
}
_PRONOUNS = re.compile(r"\b(he|him|his|she|her|hers|himself|herself)\b", re.IGNORECASE)


def _replace_pronoun(match: re.Match) -> str:
    return _PRONOUN_MAP.get(match.group(1).lower(), "they")


def _looks_like_standalone_name(line: str) -> bool:
    """Return True only for likely standalone personal names in the CV header.

    This deliberately avoids global name matching, because CVs contain many useful
    Title Case phrases that look name-like but are actually skills, awards, job
    titles, activities, or section headings.
    """
    cleaned = re.sub(r"^[\s▪•\-:]+|[\s:]+$", "", line).strip()
    lowered = re.sub(r"\s+", " ", cleaned.lower())

    if not cleaned or lowered in _HEADER_STOPWORDS:
        return False
    if any(ch.isdigit() for ch in cleaned):
        return False
    if any(sym in cleaned for sym in ["@", "/", "|", "(", ")", ":", ",", "."]):
        return False

    words = cleaned.split()
    if not (2 <= len(words) <= 4):
        return False

    word_lowers = {re.sub(r"[^a-z]", "", w.lower()) for w in words}
    if word_lowers & _NON_NAME_WORDS:
        return False

    # Avoid phrases such as "Staff Design Verification", "Physical Design Engineer".
    job_or_skill_words = {
        "engineer", "verification", "design", "physical", "staff", "senior", "junior",
        "leader", "manager", "developer", "analyst", "scientist", "consultant",
        "server", "assistant", "technician", "specialist", "intern", "trainee",
    }
    if word_lowers & job_or_skill_words:
        return False

    # Names usually have title case or all caps words, not random lowercase phrases.
    return all(re.match(r"^[A-Z][A-Za-z'\-]{1,25}$", w) or re.match(r"^[A-Z]{2,25}$", w) for w in words)


def _redact_header_names(text: str, max_lines: int = 14) -> str:
    """Redact likely candidate names in the first few lines only."""
    lines = text.splitlines()
    limit = min(max_lines, len(lines))
    previous_was_contact_heading = False

    for i in range(limit):
        raw_line = lines[i]
        stripped = raw_line.strip()
        lowered = re.sub(r"\s+", " ", stripped.lower())

        if lowered in {"contact", "coordonnées"}:
            previous_was_contact_heading = True
            continue

        if _looks_like_standalone_name(stripped):
            # Redact if it is very near the top, or directly under a Contact heading.
            if i <= 4 or previous_was_contact_heading:
                leading = re.match(r"^\s*", raw_line).group(0)
                lines[i] = leading + "[Candidate]"
                previous_was_contact_heading = False
                continue

        # Keep the contact-heading flag for one meaningful line only.
        if stripped:
            previous_was_contact_heading = False

    return "\n".join(lines)


def anonymise_text(text: str) -> str:
    """Strip PII from a CV so the LLM scores on skills and experience.

    This version is intentionally conservative: it removes clear PII while keeping
    job-relevant content such as degrees, technical skills, awards, job titles,
    company names, tools, and project descriptions.
    """
    if not isinstance(text, str):
        text = str(text)

    # 0. Remove repeated document classification labels
    text = _DOCUMENT_LABEL.sub("", text)

    # 1. Contact info and IDs
    text = _EMAIL.sub("[email]", text)
    text = _URL.sub("[url]", text)
    text = _PHONE.sub("[phone]", text)
    text = _NRIC.sub("[id]", text)

    # 2. Addresses
    text = _SG_ADDR.sub("[address]", text)
    text = _STREET_ADDR.sub("[address]", text)
    text = _SG_POSTAL.sub("[postal]", text)
    text = _CONTACT_DETAIL_LINE.sub("[Contact details redacted]", text)
    text = _ADDRESS_DETAIL_LINE.sub("[Address redacted]", text)

    # 3. Labelled personal fields and emergency-contact fields
    text = _LABELLED_NAME.sub(lambda m: m.group(1) + "[Candidate]", text)
    text = _RELATIONSHIP.sub(lambda m: m.group(1) + "[relationship]", text)
    text = _CONTACT_NO.sub(lambda m: m.group(1) + "[phone]", text)
    text = _NAME_WITH_TITLE.sub("[Candidate]", text)

    # 4. Candidate name in CV header only, not the whole document
    text = _redact_header_names(text)

    # 5. Institutions and demographics
    text = _INSTITUTION.sub("[Institution]", text)
    text = _DOB.sub("[dob]", text)
    text = _MARITAL.sub("[marital-status]", text)
    text = _LANGUAGE.sub("[language]", text)
    text = _NATIONALITY.sub("[nationality]", text)

    # 6. Pronouns: replace with neutral forms instead of deleting meaning
    text = _PRONOUNS.sub(_replace_pronoun, text)

    return text
# ── INPUT VALIDATION ──────────────────────────────────────────────────────────
MIN_CV_CHARS  = 200   # anything shorter is probably a blank/corrupt file
MIN_JD_CHARS  = 100
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 MB


def validate_uploads(jd_uploads, cv_uploads):
    """Return a list of warning strings. Empty list = all good."""
    warnings = []

    # Check for empty upload slots
    if not jd_uploads:
        warnings.append("No Job Description files uploaded.")
    if not cv_uploads:
        warnings.append("No CV files uploaded.")
    if not jd_uploads or not cv_uploads:
        return warnings  # no point checking further

    # Duplicate filenames within each slot
    jd_names = [f.name for f in jd_uploads]
    cv_names = [f.name for f in cv_uploads]
    for name in set(n for n in jd_names if jd_names.count(n) > 1):
        warnings.append(f"Duplicate JD file: '{name}' uploaded more than once.")
    for name in set(n for n in cv_names if cv_names.count(n) > 1):
        warnings.append(f"Duplicate CV file: '{name}' uploaded more than once.")

    # Same filename uploaded in both slots (JD and CV)
    cross = set(jd_names) & set(cv_names)
    for name in cross:
        warnings.append(f"'{name}' was uploaded as both a JD and a CV — is that intentional?")

    # File size
    for f in jd_uploads + cv_uploads:
        size = f.size if hasattr(f, "size") else len(f.getbuffer())
        if size == 0:
            warnings.append(f"'{f.name}' appears to be empty (0 bytes).")
        elif size > MAX_FILE_SIZE:
            warnings.append(f"'{f.name}' is larger than 5 MB — consider splitting it.")

    return warnings




def _extract_recommendation_for_share(text: str) -> str:
    if not text:
        return ""
    clean = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL | re.IGNORECASE)
    m = re.search(r"Recommendation\*?\*?\s*[:\-–]?\s*(Advance|Hold|Reject)", clean, re.IGNORECASE)
    return m.group(1).capitalize() if m else ""


# ── SHARED DEPARTMENT RESULTS ────────────────────────────────────────────────
# Prototype storage for HR-to-department handoff. On Render free tier, local
# files are suitable for demo sessions but may reset after redeploy/restart.
SHARED_RESULTS_FILE = "department_results.json"


def _safe_clean_report(text: str) -> str:
    if not isinstance(text, str):
        return ""
    return re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL | re.IGNORECASE).strip()


def make_department_safe_results(results: Dict[str, List[Dict[str, Any]]]) -> Dict[str, List[Dict[str, Any]]]:
    """Return a department-safe copy of screening results.

    Real CV filenames are replaced with Candidate 1, Candidate 2, etc. Only the
    screening report, score, summary, skills, gaps and retrieved RAG evidence are
    kept. Raw CV text/contact details are never stored here.
    """
    if not results:
        return {}

    alias_map = {}
    counter = 1
    safe_results = {}
    for jd_name, rows in results.items():
        safe_jd_name = Path(jd_name).stem.replace("_", " ").replace("-", " ")
        safe_rows = []
        for row in rows:
            real_name = row.get("cv_name", "Candidate")
            if real_name not in alias_map:
                alias_map[real_name] = f"Candidate {counter}"
                counter += 1
            safe_row = {
                "cv_name": alias_map[real_name],
                "score": row.get("score"),
                "output": _safe_clean_report(row.get("output", "")),
                "summary": row.get("summary", ""),
                "skills": row.get("skills", ""),
                "gaps": row.get("gaps", ""),
                "recommendation": _extract_recommendation_for_share(row.get("output", "")),
                "mode": row.get("mode", ""),
                "rag_context": row.get("rag_context", ""),
            }
            safe_rows.append(safe_row)
        safe_results[safe_jd_name] = safe_rows
    return safe_results


def save_department_results(results: Dict[str, List[Dict[str, Any]]], path: str = SHARED_RESULTS_FILE) -> Dict[str, List[Dict[str, Any]]]:
    safe_results = make_department_safe_results(results)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(safe_results, f, ensure_ascii=False, indent=2)
    return safe_results


def load_department_results(path: str = SHARED_RESULTS_FILE) -> Dict[str, List[Dict[str, Any]]]:
    if not os.path.exists(path):
        return {}
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def has_shared_department_results(path: str = SHARED_RESULTS_FILE) -> bool:
    return os.path.exists(path) and os.path.getsize(path) > 0

def screen_candidates(
    jd_files: Dict[str, str],
    cv_files: Dict[str, str],
    groq_key: str,
    use_rag: bool = True,
    anonymise: bool = False,
    delay: int = 1,
    progress_callback=None,
    raw_cv_files: Dict[str, str] = None,
) -> Dict[str, List[Dict[str, Any]]]:
    if not groq_key or not groq_key.strip():
        raise ValueError("Groq API key is required.")
    if not jd_files:
        raise ValueError("No job descriptions loaded.")
    if not cv_files:
        raise ValueError("No CV files loaded.")

    groq_client = Groq(api_key=groq_key.strip())

    # Apply anonymisation to CV text before chunking/screening if requested.
    # IMPORTANT: anonymise_text() must run on the RAW (un-preprocessed) text so
    # that capitalisation, @ signs and phone separators are still present.
    # raw_cv_files should be the original extracted text before preprocess_text().
    if anonymise and raw_cv_files:
        processed_cv_files = {
            name: preprocess_text(anonymise_text(raw_cv_files[name]))
            for name in cv_files
            if name in raw_cv_files
        }
    elif anonymise:
        # Fallback: anonymise already-preprocessed text (less effective but safe)
        processed_cv_files = {name: anonymise_text(text) for name, text in cv_files.items()}
    else:
        processed_cv_files = cv_files

    chunks = _build_chunks(processed_cv_files) if use_rag else []
    emb = get_embedding_model() if use_rag else None

    total_pairs = max(len(jd_files) * len(processed_cv_files), 1)
    step = 0
    all_results = {}

    for jd_name, cleaned_jd in jd_files.items():
        jd_res = []
        for cv_name, cleaned_cv in processed_cv_files.items():
            step += 1
            if progress_callback:
                progress_callback(step, total_pairs, jd_name, cv_name)

            if use_rag:
                ctx = _get_rag_context(chunks, emb, cv_name, cleaned_jd[:500])
                prompt = build_rag_prompt(cleaned_jd, cv_name, ctx)
            else:
                prompt = build_user_prompt(cleaned_jd, cv_name, cleaned_cv)

            try:
                output = _call_llm(groq_client, prompt)
            except Exception as e:
                output = f"Error: {e}"

            jd_res.append({
                "cv_name": cv_name,
                "score": extract_score(output),
                "output": output,
                "summary": extract_section(output, ["Candidate Summary"]),
                "skills": extract_section(output, ["Matching Skills"]),
                "gaps": extract_section(output, ["Missing / Weak Areas", "Missing Areas", "Weak Areas"]),
                "mode": ("RAG" if use_rag else "Direct") + (" · Anon" if anonymise else ""),
                "rag_context": ctx if use_rag else "",
            })
            if delay:
                time.sleep(delay)

        jd_res.sort(key=lambda x: x["score"] if x["score"] is not None else -1, reverse=True)
        all_results[jd_name] = jd_res

    return all_results


def build_chat_context(raw_jd_texts, raw_cv_texts, results) -> str:
    """Build a context string kept well under ~2 000 tokens so the full chat
    request (system + context + history + user question) stays inside Groq's
    free-tier per-request limit (~6 000 TPM / ~32 000 token context window).
    Rough budget: context ≤ 1 800 tokens ≈ 7 200 characters."""

    MAX_SUMMARY = 120   # chars per field
    MAX_SKILLS  = 100
    MAX_GAPS    = 100

    ctx = ""

    if raw_jd_texts:
        ctx += "=== JOB DESCRIPTIONS LOADED ===\n"
        for name in raw_jd_texts:
            ctx += f"  - {name}\n"

    if raw_cv_texts:
        ctx += "\n=== CANDIDATE CVs LOADED ===\n"
        for name in raw_cv_texts:
            ctx += f"  - {name}\n"

    if results:
        ctx += "\n=== SCREENING RESULTS ===\n"
        for jd_name, rows in results.items():
            ctx += f"\nJD: {jd_name}\n"
            for row in rows:
                score_str = f"{row['score']}/10" if row['score'] is not None else "N/A"
                verdict   = verdict_label(row['score'])[0]
                ctx += (
                    f"  [{score_str} | {verdict}] "
                    f"{Path(row['cv_name']).stem}\n"
                )
                if row.get("summary"):
                    ctx += f"    Summary: {(row['summary'] or '')[:MAX_SUMMARY]}\n"
                if row.get("skills"):
                    ctx += f"    Skills:  {(row['skills']  or '')[:MAX_SKILLS]}\n"
                if row.get("gaps"):
                    ctx += f"    Gaps:    {(row['gaps']    or '')[:MAX_GAPS]}\n"

    return ctx.strip()


def generate_shortlist_email(
    groq_key: str,
    candidate_name: str,
    job_title: str,
    company_name: str = "our company",
    extra_notes: str = "",
) -> str:
    """Generate a professional interview invitation email for a shortlisted candidate."""
    if not groq_key or not groq_key.strip():
        raise ValueError("Groq API key is required.")

    client = Groq(api_key=groq_key.strip())
    prompt = (
        f"Write a professional, warm interview invitation email to a candidate.\n\n"
        f"Candidate name: {candidate_name}\n"
        f"Job title: {job_title}\n"
        f"Company: {company_name}\n"
        f"Extra notes for personalisation: {extra_notes or 'None'}\n\n"
        f"The email should:\n"
        f"- Have a subject line on the first line starting with 'Subject: '\n"
        f"- Be concise (under 200 words)\n"
        f"- Sound human and warm, not robotic\n"
        f"- Ask them to reply with their availability for an interview\n"
        f"- Leave placeholders like [Interviewer Name] and [Company Email] where needed\n"
        f"- NOT include any commentary before or after the email\n"
        f" /nothink"
    )

    response = client.chat.completions.create(
        model=QWEN_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4,
        max_tokens=400,
    )
    raw = response.choices[0].message.content
    return re.sub(r"<think>.*?</think>", "", raw, flags=re.DOTALL | re.IGNORECASE).strip()

def ask_about_candidates(
    groq_key: str,
    context: str,
    question: str,
    chat_history: list = None,
) -> str:
    """Send a chat question to the LLM, keeping the total request size safe for
    Groq's free tier. Strategy:
      - Context is already capped by build_chat_context (~1 800 tokens).
      - We keep only the last 6 turns of chat history to avoid runaway growth.
      - max_tokens is set to 1 024 — enough for a thorough answer while leaving
        room in the 6 000 TPM budget for the prompt itself.
      - /nothink suffix disables Qwen-3's hidden chain-of-thought, preventing
        hundreds of extra tokens from being generated silently.
    """
    if not groq_key or not groq_key.strip():
        raise ValueError("Groq API key is required.")

    client = Groq(api_key=groq_key.strip())

    messages = [{"role": "system", "content": _CHAT_SYSTEM_PROMPT}]
    if context:
        messages.append({"role": "system", "content": f"Context:\n{context}"})

    # Keep only the last 6 exchanges (12 messages) to cap history size
    recent_history = (chat_history or [])[-12:]
    for item in recent_history:
        role = item.get("role")
        content = item.get("content", "")
        if role in {"user", "assistant"} and content:
            messages.append({"role": role, "content": content})

    # Append /nothink to suppress hidden reasoning tokens on Qwen-3
    messages.append({"role": "user", "content": question.rstrip() + " /nothink"})

    response = client.chat.completions.create(
        model=QWEN_MODEL,
        messages=messages,
        temperature=0.2,
        max_tokens=1024,
    )
    raw = response.choices[0].message.content
    # Strip any <think>...</think> blocks just in case
    clean = re.sub(r"<think>.*?</think>", "", raw, flags=re.DOTALL | re.IGNORECASE).strip()
    return clean
